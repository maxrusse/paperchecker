# PaperChecker v2
# Goal: Fill the provided Excel template (do not change structure) from PDFs, using cleaner LLM sub-tasks.
#
# Key improvements vs script.py:
# - Correct Excel column mapping for the provided template.
# - Handles multi-line headers (Included Articles uses 3 header rows).
# - Does NOT force the verifier to review every null cell; only non-null decisions are verified.
# - Separates extraction into smaller, less-messy LLM tasks (metadata/design, population, drugs/indication, intervention/outcomes, appraisal).
# - Writes values that match the Excel template data-validations (e.g., "Yes/No/Unclear/Not Applicable", "+1/-1/0" where expected).
#
# Dependencies:
#   pip install -U openai google-genai pymupdf python-docx openpyxl jsonschema
#
# Notes:
# - Configure OPENAI_API_KEY and GOOGLE_API_KEY via env vars or pass to run_pipeline().
# - The template is expected to already exist at TEMPLATE_XLSX.
#
# Caveat:
# - "Level of Evidence" and "Grade of Recommendation" REQUIRE a locally agreed framework (e.g., Oxford/SIGN/GRADE/AWMF).
#   This script can fill them if you provide definitions, but by default it leaves them null unless explicitly stated in the paper.

import os, json, re, copy, random, time
import urllib.parse
import urllib.request
from typing import Optional
from datetime import datetime, UTC
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz  # PyMuPDF
import openpyxl
from docx import Document

from openai import OpenAI
from google import genai
from google.genai import types
from paperchecker_utils import (
    dedupe_decisions,
    extract_page_from_evidence,
    json_pointer_get,
    normalize_excel_value,
    normalize_pmid,
    values_match,
)


# -------------------------
# CONFIG
# -------------------------
PDF_PATHS = [
    # Example:
    # "papers/paper1.pdf",
]

TEMPLATE_XLSX = None  # Auto-generated from EXCEL_MAP structure (set path to use custom template)
OUT_XLSX = f"output/mronj_extraction_{datetime.now(UTC).strftime('%Y%m%d_%H%M%S')}.xlsx"
OUT_DOCX = f"output/mronj_review_log_{datetime.now(UTC).strftime('%Y%m%d_%H%M%S')}.docx"

# Models (keep as placeholders; set to models you have access to)
OPENAI_EXTRACT_MODEL = "gpt-5.2"
OPENAI_VERIFIER_MODEL = "gpt-5.2"
GEMINI_EXTRACT_MODEL = "gemini-3-pro-preview"
GEMINI_VERIFIER_MODEL = "gemini-3-pro-preview"

REASONING_EFFORT_OPENAI = "medium"   # none|low|medium|high|xhigh
THINKING_LEVEL_GEMINI = "low"        # minimal|low|high

# Verifier-specific reasoning (lower = faster, still accurate for verification)
VERIFIER_REASONING_EFFORT_OPENAI = "low"  # none|low|medium|high|xhigh
VERIFIER_THINKING_LEVEL_GEMINI = "low"  # low|medium|high (minimal not supported)

MAX_VIEW_CHARS = 500000  # Increased - modern LLMs support large context
TASK_VIEW_CHARS = 500000  # Increased - modern LLMs support large context
VERIFIER_CHUNK_SIZE = 24
VERIFIER_PARALLEL_WORKERS = 3  # Number of parallel verification chunks
LLM_MAX_RETRIES = 3
LLM_BACKOFF_SECONDS = 2.0
LLM_BACKOFF_JITTER = 0.25

# Optional PMID lookup (PubMed E-utilities). When enabled, missing PMIDs can be
# resolved from DOI/title using a lightweight web request.
ENABLE_PUBMED_LOOKUP = True
PUBMED_API_KEY = os.getenv("PUBMED_API_KEY")
PUBMED_EMAIL = os.getenv("PUBMED_EMAIL")
PUBMED_LOOKUP_TIMEOUT = 10


# -------------------------
# EXCEL MAP (template-specific, corrected)
# -------------------------
EXCEL_MAP = {
    "sheet_key_to_name": {
        "included_articles": "Included Articles",
        "level_of_evidence": "Level of Evidence",
        "rct_appraisal": "Critical Appraisal of RCTS",
        "cohort_appraisal": "Critical Appraisal of Cohort",
        "case_series_appraisal": "Critical Appraisal of Case Seri",
        "case_control_appraisal": "Critical Appraisal of Case Cont",
        "systematic_appraisal": "Critical Appraisal of Systemati",
    },
    "sheets": {
        "included_articles": {
            "header_rows": 3,
            "key": {"field": "pmid", "col": "A"},
            "columns": {
                "pmid": "A",
                "author": "B",
                "year": "C",
                "study_design": "D",
                "n_pts": "E",
                "age_mean_years": "F",
                "gender_male_n": "G",
                "gender_female_n": "H",
                "site_maxilla": "I",
                "site_mandible": "J",
                "site_both": "K",
                "primary_cause_breast_cancer": "L",
                "primary_cause_prostate_cancer": "M",
                "primary_cause_mm": "N",
                "primary_cause_osteoporosis": "O",
                "primary_cause_other": "P",
                "ards_bisphosphonates_zoledronate": "Q",
                "ards_bisphosphonates_pamidronate": "R",
                "ards_bisphosphonates_risedronate": "S",
                "ards_bisphosphonates_alendronate": "T",
                "ards_bisphosphonates_ibandronate": "U",
                "ards_bisphosphonates_combination": "V",
                "ards_bisphosphonates_etidronate": "W",
                "ards_bisphosphonates_clodronate": "X",
                "ards_bisphosphonates_unknown_other": "Y",
                "ards_denosumab": "Z",
                "ards_both": "AA",
                "route_iv": "AB",
                "route_oral": "AC",
                "route_im": "AD",
                "route_subcutaneous": "AE",
                "route_both": "AF",
                "route_not_reported": "AG",
                "mronj_stage_at_risk": "AH",
                "mronj_stage_0": "AI",
                "prevention_technique": "AJ",
                "group_intervention": "AK",
                "group_control": "AL",
                "follow_up_mean_months": "AM",
                "follow_up_range": "AN",
                "outcome_variable": "AO",
                "mronj_development": "AP",
                "mronj_development_details": "AQ",
            },
        },
        "level_of_evidence": {
            "header_rows": 1,
            "key": {"field": "pmid", "col": "A"},
            "columns": {
                "pmid": "A",
                "author": "B",
                "year": "C",
                "study_design": "D",
                "level_of_evidence": "E",
                "grade_of_recommendation": "F",
            },
        },
        "rct_appraisal": {
            # Header row 1 is column titles, row 3 contains scoring instructions; first data row is 4.
            "header_rows": 3,
            "key": {"field": "pmid", "col": "A"},
            "columns": {
                "pmid": "A",
                "author": "B",
                "year": "C",
                "study_design": "D",
                "q1_randomized": "E",                 # "0" or "1"
                "q2_randomization_method": "F",       # "-1" | "0" | "+1"
                "q3_double_blind": "G",               # "0" or "1"
                "q4_blinding_method": "H",            # "-1" | "0" | "+1"
                "q5_withdrawals_dropouts": "I",       # "0" or "1"
                "total_score": "J",                   # integer
            },
        },
        "cohort_appraisal": {
            "header_rows": 2,
            "key": {"field": "pmid", "col": "A"},
            "columns": {
                "pmid": "A",
                "author": "B",
                "year": "C",
                "study_design": "D",
                "q1_groups_similar": "E",
                "q2_exposures_measured_similarly": "F",
                "q3_exposure_valid_reliable": "G",
                "q4_confounders_identified": "H",
                "q5_confounders_addressed": "I",
                "q6_free_of_outcome_at_start": "J",
                "q7_outcomes_valid_reliable": "K",
                "q8_followup_sufficient": "L",
                "q9_followup_complete": "M",
                "q10_address_incomplete_followup": "N",
                "q11_appropriate_statistics": "O",
            },
        },
        "case_series_appraisal": {
            "header_rows": 2,
            "key": {"field": "pmid", "col": "A"},
            "columns": {
                "pmid": "A",
                "author": "B",
                "year": "C",
                "study_design": "D",
                "q1_inclusion_criteria_clear": "E",
                "q2_condition_measured_standard": "F",
                "q3_valid_identification_methods": "G",
                "q4_consecutive_inclusion": "H",
                "q5_complete_inclusion": "I",
                "q6_demographics_reported": "J",
                "q7_clinical_info_reported": "K",
                "q8_outcomes_followup_reported": "L",
                "q9_presenting_site_reported": "M",
                "q10_statistics_appropriate": "N",
                "total_score": "O",
            },
        },
        "case_control_appraisal": {
            "header_rows": 2,
            "key": {"field": "pmid", "col": "A"},
            "columns": {
                "pmid": "A",
                "author": "B",
                "year": "C",
                "study_design": "D",
                "q1_groups_comparable": "E",
                "q2_matched_appropriately": "F",
                "q3_same_criteria_cases_controls": "G",
                "q4_exposure_valid_reliable": "H",
                "q5_exposure_measured_same_way": "I",
                "q6_confounders_identified": "J",
                "q7_confounders_addressed": "K",
                "q8_outcomes_assessed_standard": "L",
                "q9_exposure_period_long_enough": "M",
                "q10_appropriate_statistics": "N",
            },
        },
        "systematic_appraisal": {
            "header_rows": 1,
            "key": {"field": "pmid", "col": "A"},
            "columns": {
                "pmid": "A",
                "author": "B",
                "year": "C",
                "study_design": "D",
                "q1_pico": "E",
                "q2_protocol_predefined": "F",
                "q3_designs_explained": "G",
                "q4_6_search_and_duplicates": "H",
                "q7_excluded_list": "I",
                "q8_included_described": "J",
                "q9_risk_of_bias": "K",
                "q10_funding_sources": "L",
                "q11_meta_analysis_methods": "M",
                "q12_impact_of_rob": "N",
                "q13_account_for_rob": "O",
                "q14_heterogeneity_explained": "P",
                "q15_publication_bias": "Q",
                "q16_conflicts_reported": "R",
                "total_score": "S",
            },
        },
    },
}


# -------------------------
# FIELD SETS + VALUE NORMALIZATION
# -------------------------
STUDY_DESIGN_ENUM = [
    "RCT",
    "Retrospective Cohort",
    "Prospective Cohort",
    "Case-Control",
    "Retrospective Case-Series",
    "Prospective Case Series",
    "Systematic Review",
    "Metaanalysis",
]

APPRAISAL_YNUA_ENUM = ["Yes", "No", "Unclear", "Not Applicable"]
MRONJ_DEV_ENUM = ["Yes", "No"]

# -------------------------
# JSON POINTER HELPERS + MERGE
# -------------------------

def deep_merge(a, b):
    if not isinstance(a, dict) or not isinstance(b, dict):
        return copy.deepcopy(b)
    out = copy.deepcopy(a)
    for k, v in b.items():
        if k in out and isinstance(out[k], dict) and isinstance(v, dict):
            out[k] = deep_merge(out[k], v)
        else:
            out[k] = copy.deepcopy(v)
    return out

def sanitize_for_model_input(obj):
    if not isinstance(obj, dict):
        return obj
    scrub = copy.deepcopy(obj)
    for k in list(scrub.keys()):
        if k in ("verification", "validation", "model_meta", "model"):
            scrub.pop(k, None)
    return scrub


# -------------------------
# EXCEL WRITE HELPERS (avoid overwriting demo rows unless matching PMID)
# -------------------------
def column_index_from_string(col):
    col = col.upper().strip()
    idx = 0
    for c in col:
        idx = idx * 26 + (ord(c) - ord("A") + 1)
    return idx


# Human-readable column headers
COLUMN_DISPLAY_NAMES = {
    # Common fields
    "pmid": "PMID",
    "author": "Author",
    "year": "Year",
    "study_design": "Study Design",
    # Included Articles
    "n_pts": "N (Patients)",
    "age_mean_years": "Age (Mean, Years)",
    "gender_male_n": "Male (n)",
    "gender_female_n": "Female (n)",
    "site_maxilla": "Maxilla",
    "site_mandible": "Mandible",
    "site_both": "Both Sites",
    "primary_cause_breast_cancer": "Breast Cancer",
    "primary_cause_prostate_cancer": "Prostate Cancer",
    "primary_cause_mm": "Multiple Myeloma",
    "primary_cause_osteoporosis": "Osteoporosis",
    "primary_cause_other": "Other Cause",
    "ards_bisphosphonates_zoledronate": "Zoledronate",
    "ards_bisphosphonates_pamidronate": "Pamidronate",
    "ards_bisphosphonates_risedronate": "Risedronate",
    "ards_bisphosphonates_alendronate": "Alendronate",
    "ards_bisphosphonates_ibandronate": "Ibandronate",
    "ards_bisphosphonates_combination": "BP Combination",
    "ards_bisphosphonates_etidronate": "Etidronate",
    "ards_bisphosphonates_clodronate": "Clodronate",
    "ards_bisphosphonates_unknown_other": "BP Unknown/Other",
    "ards_denosumab": "Denosumab",
    "ards_both": "BP + Denosumab",
    "route_iv": "IV",
    "route_oral": "Oral",
    "route_im": "IM",
    "route_subcutaneous": "Subcutaneous",
    "route_both": "Multiple Routes",
    "route_not_reported": "Route Not Reported",
    "mronj_stage_at_risk": "At Risk Stage",
    "mronj_stage_0": "Stage 0",
    "prevention_technique": "Prevention Technique",
    "group_intervention": "Intervention Group",
    "group_control": "Control Group",
    "follow_up_mean_months": "Follow-up (Mean, Months)",
    "follow_up_range": "Follow-up Range",
    "outcome_variable": "Outcome Variable",
    "mronj_development": "MRONJ Development",
    "mronj_development_details": "MRONJ Details",
    # Level of Evidence
    "level_of_evidence": "Level of Evidence",
    "grade_of_recommendation": "Grade of Recommendation",
    # RCT Appraisal (Jadad)
    "q1_randomized": "Q1: Randomized (0/1)",
    "q2_randomization_method": "Q2: Method Described (-1/0/+1)",
    "q3_double_blind": "Q3: Double Blind (0/1)",
    "q4_blinding_method": "Q4: Blinding Method (-1/0/+1)",
    "q5_withdrawals_dropouts": "Q5: Withdrawals (0/1)",
    "total_score": "Total Score",
    # Cohort Appraisal
    "q1_groups_similar": "Q1: Groups Similar",
    "q2_exposures_measured_similarly": "Q2: Exposures Measured Similarly",
    "q3_exposure_valid_reliable": "Q3: Exposure Valid/Reliable",
    "q4_confounders_identified": "Q4: Confounders Identified",
    "q5_confounders_addressed": "Q5: Confounders Addressed",
    "q6_free_of_outcome_at_start": "Q6: Free of Outcome at Start",
    "q7_outcomes_valid_reliable": "Q7: Outcomes Valid/Reliable",
    "q8_followup_sufficient": "Q8: Follow-up Sufficient",
    "q9_followup_complete": "Q9: Follow-up Complete",
    "q10_address_incomplete_followup": "Q10: Incomplete Follow-up Addressed",
    "q11_appropriate_statistics": "Q11: Appropriate Statistics",
    # Case Series Appraisal
    "q1_inclusion_criteria_clear": "Q1: Inclusion Criteria Clear",
    "q2_condition_measured_standard": "Q2: Condition Measured Standard",
    "q3_valid_identification_methods": "Q3: Valid Identification Methods",
    "q4_consecutive_inclusion": "Q4: Consecutive Inclusion",
    "q5_complete_inclusion": "Q5: Complete Inclusion",
    "q6_demographics_reported": "Q6: Demographics Reported",
    "q7_clinical_info_reported": "Q7: Clinical Info Reported",
    "q8_outcomes_followup_reported": "Q8: Outcomes/Follow-up Reported",
    "q9_presenting_site_reported": "Q9: Presenting Site Reported",
    "q10_statistics_appropriate": "Q10: Statistics Appropriate",
    # Case Control Appraisal
    "q1_groups_comparable": "Q1: Groups Comparable",
    "q2_matched_appropriately": "Q2: Matched Appropriately",
    "q3_same_criteria_cases_controls": "Q3: Same Criteria Cases/Controls",
    "q4_exposure_valid_reliable": "Q4: Exposure Valid/Reliable",
    "q5_exposure_measured_same_way": "Q5: Exposure Measured Same Way",
    "q6_confounders_identified": "Q6: Confounders Identified",
    "q7_confounders_addressed": "Q7: Confounders Addressed",
    "q8_outcomes_assessed_standard": "Q8: Outcomes Assessed Standard",
    "q9_exposure_period_long_enough": "Q9: Exposure Period Sufficient",
    "q10_appropriate_statistics": "Q10: Appropriate Statistics",
    # Systematic Review Appraisal (AMSTAR-2)
    "q1_pico": "Q1: PICO",
    "q2_protocol_predefined": "Q2: Protocol Predefined",
    "q3_designs_explained": "Q3: Designs Explained",
    "q4_6_search_and_duplicates": "Q4-6: Search & Duplicates",
    "q7_excluded_list": "Q7: Excluded List",
    "q8_included_described": "Q8: Included Described",
    "q9_risk_of_bias": "Q9: Risk of Bias",
    "q10_funding_sources": "Q10: Funding Sources",
    "q11_meta_analysis_methods": "Q11: Meta-analysis Methods",
    "q12_impact_of_rob": "Q12: Impact of RoB",
    "q13_account_for_rob": "Q13: Account for RoB",
    "q14_heterogeneity_explained": "Q14: Heterogeneity Explained",
    "q15_publication_bias": "Q15: Publication Bias",
    "q16_conflicts_reported": "Q16: Conflicts Reported",
}


def create_template_workbook(excel_map):
    """Generate a fresh Excel workbook with headers based on EXCEL_MAP structure."""
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    # Remove default sheet
    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    header_font = Font(bold=True)
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for sheet_key, sheet_name in excel_map.get("sheet_key_to_name", {}).items():
        ws = wb.create_sheet(title=sheet_name)
        sheet_cfg = excel_map.get("sheets", {}).get(sheet_key, {})
        header_rows = int(sheet_cfg.get("header_rows", 1))
        columns = sheet_cfg.get("columns", {})

        # Sort columns by their Excel column letter
        sorted_cols = sorted(columns.items(), key=lambda x: (len(x[1]), x[1]))

        # Write headers in the last header row (so data starts right after)
        header_row = header_rows
        for field_name, col_letter in sorted_cols:
            col_idx = column_index_from_string(col_letter)
            display_name = COLUMN_DISPLAY_NAMES.get(field_name, field_name.replace("_", " ").title())
            cell = ws.cell(row=header_row, column=col_idx, value=display_name)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align

        # Set reasonable column widths
        for field_name, col_letter in sorted_cols:
            col_idx = column_index_from_string(col_letter)
            display_name = COLUMN_DISPLAY_NAMES.get(field_name, field_name)
            # Width based on header length, with min/max bounds
            width = max(10, min(30, len(display_name) + 2))
            ws.column_dimensions[get_column_letter(col_idx)].width = width

        # Freeze header rows
        ws.freeze_panes = ws.cell(row=header_rows + 1, column=1)

    return wb


def get_or_create_template(template_path, excel_map):
    """Load existing template or create new one if not found."""
    if template_path and os.path.exists(template_path):
        return template_path
    # Generate fresh template
    generated_path = "/tmp/paperchecker_template.xlsx"
    wb = create_template_workbook(excel_map)
    wb.save(generated_path)
    return generated_path

def _row_has_any_values(ws, row_idx, start_col=1, end_col=None):
    end_col = end_col or ws.max_column
    for c in range(start_col, end_col + 1):
        cell = ws.cell(row_idx, c)
        if cell.value not in (None, ""):
            return True
        if cell.data_type == "f":
            return True
        if cell.has_style:
            return True
    return False

def _find_row_by_key(ws, key_col_letter, key_value, start_row):
    if key_value in (None, ""):
        return None
    key_col_idx = column_index_from_string(key_col_letter)
    normalized_key = normalize_pmid(key_value)
    max_row = max(ws.max_row, start_row)
    for r in range(start_row, max_row + 1):
        cell_val = ws.cell(r, key_col_idx).value
        if normalized_key is not None:
            if values_match(normalize_pmid(cell_val), normalized_key):
                return r
        elif values_match(cell_val, key_value):
            return r
    return None

def _find_first_truly_empty_row(ws, start_row, end_col=None):
    end_col = end_col or ws.max_column
    max_row = max(ws.max_row, start_row)
    for r in range(start_row, max_row + 1):
        if not _row_has_any_values(ws, r, 1, end_col):
            return r
    return max_row + 1

def _resolve_anchor_row(wb, pmid):
    inc_cfg = EXCEL_MAP["sheets"]["included_articles"]
    ws = wb[EXCEL_MAP["sheet_key_to_name"]["included_articles"]]
    start_row = int(inc_cfg["header_rows"]) + 1
    # 1) If PMID exists, use that row.
    found = _find_row_by_key(ws, inc_cfg["key"]["col"], pmid, start_row)
    if found is not None:
        return found
    # 2) Otherwise, append after the last fully empty row region (do not overwrite demo-like rows with other values).
    return _find_first_truly_empty_row(ws, start_row, end_col=ws.max_column)

def _clear_data_rows(wb, excel_map):
    """Clear all data rows (below headers) from all sheets, keeping only headers."""
    for sheet_key, sheet_name in (excel_map.get("sheet_key_to_name") or {}).items():
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        sheet_cfg = (excel_map.get("sheets") or {}).get(sheet_key) or {}
        header_rows = int(sheet_cfg.get("header_rows") or 1)
        start_row = header_rows + 1

        # Delete all rows below header
        max_row = ws.max_row
        if max_row >= start_row:
            ws.delete_rows(start_row, max_row - start_row + 1)


def apply_to_workbook(final_obj, template_xlsx, out_xlsx, excel_map, clear_existing_data=False):
    wb = openpyxl.load_workbook(template_xlsx)

    # Clear old/demo data rows if requested (typically on first PDF only)
    if clear_existing_data:
        _clear_data_rows(wb, excel_map)

    sheets_data = ((final_obj.get("record") or {}).get("sheets")) or {}
    pmid = (final_obj.get("paper_id") or {}).get("pmid")

    # Anchor row comes from Included Articles so relevant sheets align.
    anchor_row = _resolve_anchor_row(wb, pmid)

    def write_sheet(sheet_key):
        sheet_name = (excel_map.get("sheet_key_to_name") or {}).get(sheet_key)
        if not sheet_name or sheet_name not in wb.sheetnames:
            return
        ws = wb[sheet_name]
        sheet_cfg = (excel_map.get("sheets") or {}).get(sheet_key) or {}
        header_rows = int(sheet_cfg.get("header_rows") or 1)
        start_row = header_rows + 1

        key_cfg = sheet_cfg.get("key") or {"field": "pmid", "col": "A"}
        key_col = key_cfg.get("col") or "A"
        row_idx = _find_row_by_key(ws, key_col, pmid, start_row)
        if row_idx is None:
            row_idx = max(anchor_row, start_row)

        cols = sheet_cfg.get("columns") or {}
        payload = sheets_data.get(sheet_key)

        # Write PMID only for sheets we actually populate (or always for Included Articles / Level of Evidence).
        if "pmid" in cols:
            ws[f"{cols['pmid']}{row_idx}"].value = pmid

        if isinstance(payload, dict):
            for field, col_letter in cols.items():
                if field == "pmid":
                    continue
                if field in payload:
                    v = payload.get(field)
                    if v is None:
                        continue
                    ws[f"{col_letter}{row_idx}"].value = normalize_excel_value(v)

        # Back-fill author/year/study_design from Included Articles if available and blank.
        inc = sheets_data.get("included_articles") or {}
        if isinstance(inc, dict):
            for f in ("author", "year", "study_design"):
                if f in cols and ws[f"{cols[f]}{row_idx}"].value in (None, ""):
                    if inc.get(f) not in (None, ""):
                        ws[f"{cols[f]}{row_idx}"].value = normalize_excel_value(inc.get(f))

    # Always write these two sheets (even if sparse): boss deliverable tables.
    write_sheet("included_articles")
    write_sheet("level_of_evidence")

    # Only write the appraisal sheet(s) that actually exist as dict payloads.
    for sheet_key in ("rct_appraisal", "cohort_appraisal", "case_series_appraisal", "case_control_appraisal", "systematic_appraisal"):
        if isinstance(sheets_data.get(sheet_key), dict):
            write_sheet(sheet_key)

    wb.save(out_xlsx)



# -------------------------
# SCORE COMPUTATION (template-aligned)
# -------------------------
def _is_yes(v):
    if v is None:
        return False
    if isinstance(v, str):
        return v.strip() == "Yes"
    return False

def compute_scores_inplace(final_obj):
    sheets = (final_obj.get("record") or {}).get("sheets") or {}

    # Jadad-like scoring for RCT sheet.
    rct = sheets.get("rct_appraisal")
    if isinstance(rct, dict):
        total = 0
        for k in ("q1_randomized", "q3_double_blind", "q5_withdrawals_dropouts"):
            if str(rct.get(k)).strip() == "1":
                total += 1
        for k in ("q2_randomization_method", "q4_blinding_method"):
            vv = str(rct.get(k)).strip()
            if vv == "+1":
                total += 1
            elif vv == "-1":
                total -= 1
        rct["total_score"] = max(0, int(total))

    # Case series: count Yes across 10 questions.
    cs = sheets.get("case_series_appraisal")
    if isinstance(cs, dict):
        q_keys = [k for k in cs.keys() if k.startswith("q")]
        total = sum(1 for k in q_keys if _is_yes(cs.get(k)))
        cs["total_score"] = int(total)

    # Systematic: count Yes across q* columns (including the combined q4_6 as one column).
    sr = sheets.get("systematic_appraisal")
    if isinstance(sr, dict):
        q_keys = [k for k in sr.keys() if k.startswith("q")]
        total = sum(1 for k in q_keys if _is_yes(sr.get(k)))
        sr["total_score"] = int(total)


# -------------------------
# PDF TEXT (page-aware) + TARGETED VIEWS
# -------------------------
def extract_pdf_pages(pdf_path):
    doc = fitz.open(pdf_path)
    pages = []
    for i in range(doc.page_count):
        txt = doc.load_page(i).get_text("text")
        pages.append({"page_index": i, "text": txt or ""})
    doc.close()
    return pages

DOI_REGEX = re.compile(r"\b10\.\d{4,9}/[^\s\"<>]+", re.IGNORECASE)
TITLE_STOPWORDS = {
    "abstract",
    "introduction",
    "keywords",
    "correspondence",
    "author",
    "authors",
    "journal",
    "doi",
}

def _normalize_doi(doi: str) -> str:
    doi = doi.strip()
    doi = re.sub(r"^(doi\s*:\s*)", "", doi, flags=re.IGNORECASE)
    return doi.rstrip(").,;")

def _extract_doi(text: str) -> Optional[str]:
    if not text:
        return None
    match = DOI_REGEX.search(text)
    if not match:
        return None
    return _normalize_doi(match.group(0))

def _extract_title_from_page(text: str) -> Optional[str]:
    if not text:
        return None
    candidates = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        lower = line.lower()
        if any(stop in lower for stop in TITLE_STOPWORDS):
            continue
        if sum(ch.isdigit() for ch in line) > 3:
            continue
        words = line.split()
        if not (5 <= len(words) <= 25):
            continue
        candidates.append((len(words), len(line), line))
    if not candidates:
        return None
    candidates.sort(reverse=True)
    return candidates[0][2]

def extract_paper_id_from_pages(pages: list[dict]) -> dict:
    if not pages:
        return {"doi": None, "title": None}
    first_pages = pages[:2]
    combined = "\n".join(p.get("text", "") for p in first_pages)
    doi = _extract_doi(combined)
    title = _extract_title_from_page(first_pages[0].get("text", ""))
    return {"doi": doi, "title": title}

def _clean_text(t):
    t = re.sub(r"[ \t]+\n", "\n", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t

def make_global_view(pages, max_chars=MAX_VIEW_CHARS):
    # A compact view: first pages + windows around common section headers.
    full = "\n".join([p["text"] for p in pages])
    full = _clean_text(full)
    tl = full.lower()

    def win_at(needle, span=12000):
        idx = tl.find(needle)
        if idx == -1:
            return ""
        start = max(0, idx - 1500)
        end = min(len(full), idx + span)
        return full[start:end]

    chunks = []
    chunks.append(full[:8000])
    for key in ["abstract", "methods", "materials and methods", "results", "discussion", "conclusion", "table"]:
        c = win_at(key)
        if c:
            chunks.append("\n\n===== " + key.upper() + " (WINDOW) =====\n" + c)

    combined = "\n".join(chunks)
    return combined[:max_chars]

def make_task_view(pages, keywords, max_chars=TASK_VIEW_CHARS, window=1400):
    # Lightweight retrieval: collect small windows from pages containing any keyword.
    kws = [k.lower() for k in keywords if k]
    out = []
    for p in pages:
        tx = p["text"] or ""
        tl = tx.lower()
        hits = [k for k in kws if k in tl]
        if not hits:
            continue
        # Take multiple windows per page around the first few hits.
        # This is intentionally simple; it dramatically reduces LLM confusion vs a single huge prompt.
        for k in hits[:4]:
            idx = tl.find(k)
            if idx == -1:
                continue
            start = max(0, idx - window)
            end = min(len(tx), idx + window)
            snippet = tx[start:end]
            snippet = _clean_text(snippet)
            out.append(f"\n\n--- PAGE {p['page_index'] + 1} (hit: {k}) ---\n{snippet}")
        if sum(len(s) for s in out) >= max_chars:
            break
    joined = "\n".join(out)
    if len(joined) < 1500:
        # Fallback: global view
        return make_global_view(pages, max_chars=max_chars)
    return joined[:max_chars]

def _select_pages(pages, page_numbers):
    if not pages:
        return []
    if not page_numbers:
        return pages
    page_set = set(page_numbers)
    subset = [p for p in pages if p.get("page_index") in page_set]
    return subset or pages


DECISION_KEYWORDS = {
    "pmid": ["pmid"],
    "doi": ["doi"],
    "title": ["title"],
    "author": ["author"],
    "year": ["year"],
    "study_design": ["study design", "randomized", "cohort", "case", "systematic review"],
    "n_pts": ["participants", "patients", "sample", "n="],
    "age_mean_years": ["mean age", "age"],
    "gender_male_n": ["male", "men"],
    "gender_female_n": ["female", "women"],
    "prevention_technique": ["prevention", "technique"],
    "group_intervention": ["intervention", "treatment"],
    "group_control": ["control", "comparison"],
    "follow_up_mean_months": ["follow-up", "months"],
    "follow_up_range": ["follow-up", "range"],
    "outcome_variable": ["outcome", "endpoint"],
    "mronj_development": ["mronj", "osteonecrosis"],
    "mronj_development_details": ["mronj", "osteonecrosis"],
    "level_of_evidence": ["level of evidence"],
    "grade_of_recommendation": ["grade of recommendation"],
}

def decision_keywords_for_path(path):
    if not path:
        return []
    leaf = path.strip("/").split("/")[-1]
    if leaf.startswith("q"):
        return ["methods", "random", "blind", "confound", "follow-up", "risk of bias"]
    return DECISION_KEYWORDS.get(leaf, [])

def build_verifier_view(pages, decisions):
    keywords = []
    for d in decisions or []:
        keywords.extend(decision_keywords_for_path(d.get("path")))
    keywords = [k for k in dict.fromkeys(keywords) if k]
    if not keywords:
        return make_global_view(pages, max_chars=MAX_VIEW_CHARS)
    return make_task_view(pages, keywords, max_chars=MAX_VIEW_CHARS)


# -------------------------
# JSON SCHEMAS (task outputs)
# -------------------------
def _sheet_schema_included_articles_partial():
    props = {
        "pmid": {"type": ["integer", "null"]},
        "author": {"type": ["string", "null"]},
        "year": {"type": ["integer", "null"]},
        "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
        "n_pts": {"type": ["integer", "null"]},
        "age_mean_years": {"type": ["number", "null"]},
        "gender_male_n": {"type": ["integer", "null"]},
        "gender_female_n": {"type": ["integer", "null"]},
        "site_maxilla": {"type": ["integer", "null"], "enum": [1, None]},
        "site_mandible": {"type": ["integer", "null"], "enum": [1, None]},
        "site_both": {"type": ["integer", "null"], "enum": [1, None]},
        "primary_cause_breast_cancer": {"type": ["integer", "null"], "enum": [1, None]},
        "primary_cause_prostate_cancer": {"type": ["integer", "null"], "enum": [1, None]},
        "primary_cause_mm": {"type": ["integer", "null"], "enum": [1, None]},
        "primary_cause_osteoporosis": {"type": ["integer", "null"], "enum": [1, None]},
        "primary_cause_other": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_zoledronate": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_pamidronate": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_risedronate": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_alendronate": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_ibandronate": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_combination": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_etidronate": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_clodronate": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_bisphosphonates_unknown_other": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_denosumab": {"type": ["integer", "null"], "enum": [1, None]},
        "ards_both": {"type": ["integer", "null"], "enum": [1, None]},
        "route_iv": {"type": ["integer", "null"], "enum": [1, None]},
        "route_oral": {"type": ["integer", "null"], "enum": [1, None]},
        "route_im": {"type": ["integer", "null"], "enum": [1, None]},
        "route_subcutaneous": {"type": ["integer", "null"], "enum": [1, None]},
        "route_both": {"type": ["integer", "null"], "enum": [1, None]},
        "route_not_reported": {"type": ["integer", "null"], "enum": [1, None]},
        "mronj_stage_at_risk": {"type": ["integer", "null"], "enum": [1, None]},
        "mronj_stage_0": {"type": ["integer", "null"], "enum": [1, None]},
        "prevention_technique": {"type": ["string", "null"]},
        "group_intervention": {"type": ["string", "null"]},
        "group_control": {"type": ["string", "null"]},
        "follow_up_mean_months": {"type": ["number", "null"]},
        "follow_up_range": {"type": ["string", "null"]},
        "outcome_variable": {"type": ["string", "null"]},
        "mronj_development": {"type": ["string", "null"], "enum": MRONJ_DEV_ENUM + [None]},
        "mronj_development_details": {"type": ["string", "null"]},
    }
    return {"type": "object", "additionalProperties": False, "required": list(props.keys()), "properties": props}

def _sheet_schema_level_of_evidence_partial():
    props = {
        "pmid": {"type": ["integer", "null"]},
        "author": {"type": ["string", "null"]},
        "year": {"type": ["integer", "null"]},
        "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
        "level_of_evidence": {"type": ["string", "null"]},
        "grade_of_recommendation": {"type": ["string", "null"]},
    }
    return {"type": "object", "additionalProperties": False, "required": list(props.keys()), "properties": props}

DECISION_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": ["path", "value", "evidence", "page", "is_critical"],
    "properties": {
        "path": {"type": "string"},
        "value": {"type": ["string", "number", "integer", "boolean", "null"]},
        "evidence": {"type": "string"},
        "page": {"type": ["integer", "null"]},  # 1-based page number if known
        "is_critical": {"type": "boolean"},
    },
}

PAPER_ID_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": ["pmid", "doi", "title"],
    "properties": {
        "pmid": {"type": ["integer", "null"]},
        "doi": {"type": ["string", "null"]},
        "title": {"type": ["string", "null"]},
    },
}

STUDY_TYPE_ENUM = ["rct", "cohort", "case_series", "case_control", "systematic_review", "other", "unclear"]

def build_task_schema(task_name, allowed_sheet_key=None, allowed_included_keys=None, allowed_level_keys=None):
    # Schema restricts patch to only the sheet/fields for this task.
    inc_schema = _sheet_schema_included_articles_partial()
    lev_schema = _sheet_schema_level_of_evidence_partial()

    # If allowed_included_keys provided, shrink properties to only those keys.
    if allowed_included_keys is not None:
        inc_schema = copy.deepcopy(inc_schema)
        inc_schema["properties"] = {k: v for k, v in inc_schema["properties"].items() if k in allowed_included_keys}
        inc_schema["required"] = list(inc_schema["properties"].keys())

    if allowed_level_keys is not None:
        lev_schema = copy.deepcopy(lev_schema)
        lev_schema["properties"] = {k: v for k, v in lev_schema["properties"].items() if k in allowed_level_keys}
        lev_schema["required"] = list(lev_schema["properties"].keys())

    sheets_props = {}
    if allowed_sheet_key == "included_articles":
        sheets_props["included_articles"] = inc_schema
    elif allowed_sheet_key == "level_of_evidence":
        sheets_props["level_of_evidence"] = lev_schema
    else:
        # allow both by default
        sheets_props["included_articles"] = inc_schema
        sheets_props["level_of_evidence"] = lev_schema

    for sheet_schema in sheets_props.values():
        sheet_schema["required"] = list(sheet_schema.get("properties", {}).keys())

    schema = {
        "type": "object",
        "additionalProperties": False,
        "required": ["patch", "decisions", "confidence", "notes"],
        "properties": {
            "patch": {
                "type": "object",
                "additionalProperties": False,
                "required": ["paper_id", "study_type", "record"],
                "properties": {
                    "paper_id": PAPER_ID_SCHEMA,
                    "study_type": {"type": ["string", "null"], "enum": STUDY_TYPE_ENUM + [None]},
                    "record": {
                        "type": "object",
                        "additionalProperties": False,
                        "required": ["sheets"],
                        "properties": {
                            "sheets": {
                                "type": "object",
                                "additionalProperties": False,
                                "required": list(sheets_props.keys()),
                                "properties": sheets_props,
                            }
                        },
                    },
                },
            },
            "decisions": {"type": "array", "items": DECISION_SCHEMA},
            "confidence": {"type": "number", "minimum": 0.0, "maximum": 1.0},
            "notes": {"type": "string"},
        },
    }
    return schema

def build_appraisal_schema(study_type):
    # Each appraisal task schema only allows the relevant appraisal sheet keys.
    # All answers should be strings matching the Excel validation lists where present.
    def y_schema():
        return {"type": ["string", "null"], "enum": APPRAISAL_YNUA_ENUM + [None]}

    sheets = {}

    if study_type == "rct":
        sheets["rct_appraisal"] = {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "pmid": {"type": ["integer", "null"]},
                "author": {"type": ["string", "null"]},
                "year": {"type": ["integer", "null"]},
                "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
                "q1_randomized": {"type": ["string", "null"], "enum": ["0", "1", None]},
                "q2_randomization_method": {"type": ["string", "null"], "enum": ["-1", "0", "+1", None]},
                "q3_double_blind": {"type": ["string", "null"], "enum": ["0", "1", None]},
                "q4_blinding_method": {"type": ["string", "null"], "enum": ["-1", "0", "+1", None]},
                "q5_withdrawals_dropouts": {"type": ["string", "null"], "enum": ["0", "1", None]},
                "total_score": {"type": ["integer", "null"]},
            },
        }
    elif study_type == "cohort":
        sheets["cohort_appraisal"] = {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "pmid": {"type": ["integer", "null"]},
                "author": {"type": ["string", "null"]},
                "year": {"type": ["integer", "null"]},
                "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
                "q1_groups_similar": y_schema(),
                "q2_exposures_measured_similarly": y_schema(),
                "q3_exposure_valid_reliable": y_schema(),
                "q4_confounders_identified": y_schema(),
                "q5_confounders_addressed": y_schema(),
                "q6_free_of_outcome_at_start": y_schema(),
                "q7_outcomes_valid_reliable": y_schema(),
                "q8_followup_sufficient": y_schema(),
                "q9_followup_complete": y_schema(),
                "q10_address_incomplete_followup": y_schema(),
                "q11_appropriate_statistics": y_schema(),
            },
        }
    elif study_type == "case_series":
        sheets["case_series_appraisal"] = {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "pmid": {"type": ["integer", "null"]},
                "author": {"type": ["string", "null"]},
                "year": {"type": ["integer", "null"]},
                "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
                "q1_inclusion_criteria_clear": y_schema(),
                "q2_condition_measured_standard": y_schema(),
                "q3_valid_identification_methods": y_schema(),
                "q4_consecutive_inclusion": y_schema(),
                "q5_complete_inclusion": y_schema(),
                "q6_demographics_reported": y_schema(),
                "q7_clinical_info_reported": y_schema(),
                "q8_outcomes_followup_reported": y_schema(),
                "q9_presenting_site_reported": y_schema(),
                "q10_statistics_appropriate": y_schema(),
                "total_score": {"type": ["integer", "null"]},
            },
        }
    elif study_type == "case_control":
        sheets["case_control_appraisal"] = {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "pmid": {"type": ["integer", "null"]},
                "author": {"type": ["string", "null"]},
                "year": {"type": ["integer", "null"]},
                "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
                "q1_groups_comparable": y_schema(),
                "q2_matched_appropriately": y_schema(),
                "q3_same_criteria_cases_controls": y_schema(),
                "q4_exposure_valid_reliable": y_schema(),
                "q5_exposure_measured_same_way": y_schema(),
                "q6_confounders_identified": y_schema(),
                "q7_confounders_addressed": y_schema(),
                "q8_outcomes_assessed_standard": y_schema(),
                "q9_exposure_period_long_enough": y_schema(),
                "q10_appropriate_statistics": y_schema(),
            },
        }
    elif study_type == "systematic_review":
        sheets["systematic_appraisal"] = {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "pmid": {"type": ["integer", "null"]},
                "author": {"type": ["string", "null"]},
                "year": {"type": ["integer", "null"]},
                "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
                "q1_pico": y_schema(),
                "q2_protocol_predefined": y_schema(),
                "q3_designs_explained": y_schema(),
                "q4_6_search_and_duplicates": y_schema(),
                "q7_excluded_list": y_schema(),
                "q8_included_described": y_schema(),
                "q9_risk_of_bias": y_schema(),
                "q10_funding_sources": y_schema(),
                "q11_meta_analysis_methods": y_schema(),
                "q12_impact_of_rob": y_schema(),
                "q13_account_for_rob": y_schema(),
                "q14_heterogeneity_explained": y_schema(),
                "q15_publication_bias": y_schema(),
                "q16_conflicts_reported": y_schema(),
                "total_score": {"type": ["integer", "null"]},
            },
        }
    else:
        sheets = {}

    for sheet_schema in sheets.values():
        sheet_schema["required"] = list(sheet_schema.get("properties", {}).keys())

    return {
        "type": "object",
        "additionalProperties": False,
        "required": ["patch", "decisions", "confidence", "notes"],
        "properties": {
            "patch": {
                "type": "object",
                "additionalProperties": False,
                "required": ["record"],
                "properties": {
                    "record": {
                        "type": "object",
                        "additionalProperties": False,
                        "required": ["sheets"],
                        "properties": {
                            "sheets": {
                                "type": "object",
                                "additionalProperties": False,
                                "required": list(sheets.keys()),
                                "properties": sheets,
                            }
                        },
                    },
                },
            },
            "decisions": {"type": "array", "items": DECISION_SCHEMA},
            "confidence": {"type": "number", "minimum": 0.0, "maximum": 1.0},
            "notes": {"type": "string"},
        },
    }


def _optionalize_object_schema(s: dict) -> dict:
    """Create an optional version of an object schema (required: all properties)."""
    s2 = copy.deepcopy(s)
    s2["required"] = list(s2.get("properties", {}).keys())
    return s2


def _suggested_patch_schema():
    """
    suggested_patch is ALWAYS an object (never null) so we avoid object|null unions
    that can break OpenAI's strict schema validation.

    It is "deeply optional": required lists are empty at every level, so the model
    can return {} when no corrections are needed, or only the corrected fields.
    """
    def y_schema():
        return {"type": ["string", "null"], "enum": APPRAISAL_YNUA_ENUM + [None]}

    paper_id_obj = {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {
            "pmid": {"type": ["integer", "null"]},
            "doi": {"type": ["string", "null"]},
            "title": {"type": ["string", "null"]},
        },
    }

    included_articles_obj = copy.deepcopy(_sheet_schema_included_articles_partial())
    included_articles_obj["required"] = []

    level_of_evidence_obj = copy.deepcopy(_sheet_schema_level_of_evidence_partial())
    level_of_evidence_obj["required"] = []

    rct_obj = {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {
            "pmid": {"type": ["integer", "null"]},
            "author": {"type": ["string", "null"]},
            "year": {"type": ["integer", "null"]},
            "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
            "q1_randomized": {"type": ["string", "null"], "enum": ["0", "1", None]},
            "q2_randomization_method": {"type": ["string", "null"], "enum": ["-1", "0", "+1", None]},
            "q3_double_blind": {"type": ["string", "null"], "enum": ["0", "1", None]},
            "q4_blinding_method": {"type": ["string", "null"], "enum": ["-1", "0", "+1", None]},
            "q5_withdrawals_dropouts": {"type": ["string", "null"], "enum": ["0", "1", None]},
            "total_score": {"type": ["integer", "null"]},
        },
    }

    cohort_obj = {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {
            "pmid": {"type": ["integer", "null"]},
            "author": {"type": ["string", "null"]},
            "year": {"type": ["integer", "null"]},
            "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
            "q1_groups_similar": y_schema(),
            "q2_exposures_measured_similarly": y_schema(),
            "q3_exposure_valid_reliable": y_schema(),
            "q4_confounders_identified": y_schema(),
            "q5_confounders_addressed": y_schema(),
            "q6_free_of_outcome_at_start": y_schema(),
            "q7_outcomes_valid_reliable": y_schema(),
            "q8_followup_sufficient": y_schema(),
            "q9_followup_complete": y_schema(),
            "q10_address_incomplete_followup": y_schema(),
            "q11_appropriate_statistics": y_schema(),
        },
    }

    case_series_obj = {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {
            "pmid": {"type": ["integer", "null"]},
            "author": {"type": ["string", "null"]},
            "year": {"type": ["integer", "null"]},
            "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
            "q1_inclusion_criteria_clear": y_schema(),
            "q2_condition_measured_standard": y_schema(),
            "q3_valid_identification_methods": y_schema(),
            "q4_consecutive_inclusion": y_schema(),
            "q5_complete_inclusion": y_schema(),
            "q6_demographics_reported": y_schema(),
            "q7_clinical_info_reported": y_schema(),
            "q8_outcomes_followup_reported": y_schema(),
            "q9_presenting_site_reported": y_schema(),
            "q10_statistics_appropriate": y_schema(),
            "total_score": {"type": ["integer", "null"]},
        },
    }

    case_control_obj = {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {
            "pmid": {"type": ["integer", "null"]},
            "author": {"type": ["string", "null"]},
            "year": {"type": ["integer", "null"]},
            "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
            "q1_groups_comparable": y_schema(),
            "q2_matched_appropriately": y_schema(),
            "q3_same_criteria_cases_controls": y_schema(),
            "q4_exposure_valid_reliable": y_schema(),
            "q5_exposure_measured_same_way": y_schema(),
            "q6_confounders_identified": y_schema(),
            "q7_confounders_addressed": y_schema(),
            "q8_outcomes_assessed_standard": y_schema(),
            "q9_exposure_period_long_enough": y_schema(),
            "q10_appropriate_statistics": y_schema(),
        },
    }

    systematic_obj = {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {
            "pmid": {"type": ["integer", "null"]},
            "author": {"type": ["string", "null"]},
            "year": {"type": ["integer", "null"]},
            "study_design": {"type": ["string", "null"], "enum": STUDY_DESIGN_ENUM + [None]},
            "q1_pico": y_schema(),
            "q2_protocol_predefined": y_schema(),
            "q3_designs_explained": y_schema(),
            "q4_6_search_and_duplicates": y_schema(),
            "q7_excluded_list": y_schema(),
            "q8_included_described": y_schema(),
            "q9_risk_of_bias": y_schema(),
            "q10_funding_sources": y_schema(),
            "q11_meta_analysis_methods": y_schema(),
            "q12_impact_of_rob": y_schema(),
            "q13_account_for_rob": y_schema(),
            "q14_heterogeneity_explained": y_schema(),
            "q15_publication_bias": y_schema(),
            "q16_conflicts_reported": y_schema(),
            "total_score": {"type": ["integer", "null"]},
        },
    }

    sheets_obj = {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {
            "included_articles": included_articles_obj,
            "level_of_evidence": level_of_evidence_obj,
            "rct_appraisal": rct_obj,
            "cohort_appraisal": cohort_obj,
            "case_series_appraisal": case_series_obj,
            "case_control_appraisal": case_control_obj,
            "systematic_appraisal": systematic_obj,
        },
    }

    record_obj = {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {"sheets": sheets_obj},
    }

    return {
        "type": "object",
        "additionalProperties": False,
        "required": [],
        "properties": {
            "paper_id": paper_id_obj,
            "study_type": {"type": ["string", "null"], "enum": STUDY_TYPE_ENUM + [None]},
            "record": record_obj,
        },
    }


VERIFIER_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": ["verdict", "critical_errors", "decision_reviews", "suggested_patch", "rationale", "confidence"],
    "properties": {
        "verdict": {"type": "string", "enum": ["AGREE", "DISAGREE", "UNSURE"]},
        "critical_errors": {"type": "array", "items": {"type": "string"}},
        "decision_reviews": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["path", "is_critical", "status", "driver_value", "proposed_value", "explanation", "evidence"],
                "properties": {
                    "path": {"type": "string"},
                    "is_critical": {"type": "boolean"},
                    "status": {"type": "string", "enum": ["AGREE", "DISAGREE", "UNSURE"]},
                    "driver_value": {"type": ["string", "number", "integer", "boolean", "null"]},
                    "proposed_value": {"type": ["string", "number", "integer", "boolean", "null"]},
                    "explanation": {"type": "string"},
                    "evidence": {"type": "string"},
                },
            },
        },
        "suggested_patch": _suggested_patch_schema(),
        "rationale": {"type": "string"},
        "confidence": {"type": "number", "minimum": 0.0, "maximum": 1.0},
    },
}


# -------------------------
# PROMPTS (small, focused tasks)
# -------------------------
TASK_SYSTEM = (
    "You extract structured evidence from a single MRONJ prevention paper.\n"
    "Use ONLY the provided text. Do not guess.\n"
    "If not reported, return null.\n"
    "For 1/0 flag fields: use 1 when explicitly present, otherwise null.\n"
    "Evidence must be short (1 sentence). No long quotes.\n"
    "If you can infer a page number from the snippet header (e.g., '--- PAGE 3'), include that page.\n"
    "Return strict JSON matching the schema.\n"
)

VERIFIER_SYSTEM = (
    "You are an independent verifier.\n"
    "Check whether each listed decision is supported by the provided paper text.\n"
    "For each decision: AGREE, DISAGREE (with proposed_value), or UNSURE.\n"
    "Evidence must be short (1 sentence), no long quotes.\n"
    "If DISAGREE, propose the minimal corrected value.\n"
    "Also provide suggested_patch as a minimal JSON object patch (only corrected fields).\n"
    "Return strict JSON matching the schema.\n"
)

def _task_user(task_name, allowed_fields_text, view_text, context_json=None):
    ctx = ""
    if context_json is not None:
        ctx = "\n\nCONTEXT_JSON (already extracted; do not change unrelated fields):\n" + json.dumps(context_json, ensure_ascii=True)
    return (
        f"TASK_NAME: {task_name}\n"
        f"FIELDS_TO_FILL:\n{allowed_fields_text}\n\n"
        f"PAPER_TEXT (TASK VIEW):\n{view_text}\n"
        + ctx
    )


# -------------------------
# LLM CALLS
# -------------------------
def _call_with_retries(fn, description):
    last_exc = None
    for attempt in range(1, LLM_MAX_RETRIES + 1):
        try:
            return fn()
        except Exception as exc:
            last_exc = exc
            if attempt >= LLM_MAX_RETRIES:
                break
            backoff = LLM_BACKOFF_SECONDS * (2 ** (attempt - 1))
            jitter = random.uniform(0, LLM_BACKOFF_JITTER)
            time.sleep(backoff + jitter)
    raise RuntimeError(f"{description} failed after {LLM_MAX_RETRIES} attempts: {last_exc}") from last_exc

def openai_json(oai_client, system_text, user_text, schema, schema_name):
    def _call():
        resp = oai_client.responses.create(
            model=OPENAI_EXTRACT_MODEL,
            reasoning={"effort": REASONING_EFFORT_OPENAI},
            input=[
                {"role": "system", "content": system_text},
                {"role": "user", "content": user_text},
            ],
            text={"format": {"type": "json_schema", "name": schema_name, "schema": schema, "strict": True}},
        )
        return json.loads(resp.output_text)
    return _call_with_retries(_call, f"OpenAI call ({schema_name})")

def gemini_json(gclient, system_text, user_text, schema):
    def _call():
        resp = gclient.models.generate_content(
            model=GEMINI_EXTRACT_MODEL,
            contents=user_text,
            config=types.GenerateContentConfig(
                system_instruction=system_text,
                response_mime_type="application/json",
                response_json_schema=schema,
                thinking_config=types.ThinkingConfig(thinking_level=THINKING_LEVEL_GEMINI),
                temperature=0.0,
            ),
        )
        return json.loads(resp.text)
    return _call_with_retries(_call, "Gemini call")


# Verifier-specific LLM functions (use lower reasoning for faster verification)
def openai_json_verifier(oai_client, system_text, user_text, schema, schema_name):
    def _call():
        resp = oai_client.responses.create(
            model=OPENAI_VERIFIER_MODEL,
            reasoning={"effort": VERIFIER_REASONING_EFFORT_OPENAI},
            input=[
                {"role": "system", "content": system_text},
                {"role": "user", "content": user_text},
            ],
            text={"format": {"type": "json_schema", "name": schema_name, "schema": schema, "strict": True}},
        )
        return json.loads(resp.output_text)
    return _call_with_retries(_call, f"OpenAI verifier ({schema_name})")


def gemini_json_verifier(gclient, system_text, user_text, schema):
    def _call():
        resp = gclient.models.generate_content(
            model=GEMINI_VERIFIER_MODEL,
            contents=user_text,
            config=types.GenerateContentConfig(
                system_instruction=system_text,
                response_mime_type="application/json",
                response_json_schema=schema,
                thinking_config=types.ThinkingConfig(thinking_level=VERIFIER_THINKING_LEVEL_GEMINI),
                temperature=0.0,
            ),
        )
        return json.loads(resp.text)
    return _call_with_retries(_call, "Gemini verifier")


# -------------------------
# DECISION UTILITIES
# -------------------------
def decisions_only_non_null(decisions):
    out = []
    for d in decisions or []:
        if d.get("value") is None:
            continue
        if d.get("page") is None and d.get("evidence"):
            extracted = extract_page_from_evidence(d.get("evidence"))
            if extracted is not None:
                d = copy.deepcopy(d)
                d["page"] = extracted
        out.append(d)
    return out

def chunk_list(xs, n):
    return [xs[i:i+n] for i in range(0, len(xs), n)]


def group_decisions_by_page(decisions, max_chunk_size=VERIFIER_CHUNK_SIZE):
    """Group decisions by page for more efficient verification.

    Decisions on the same pages are grouped together, then split into chunks
    that don't exceed max_chunk_size. This ensures the verifier gets focused
    context (only relevant pages) for each chunk.
    """
    # Group by page (None pages go together)
    by_page = {}
    for d in decisions:
        page = d.get("page")
        by_page.setdefault(page, []).append(d)

    # Build chunks: try to keep same-page decisions together
    chunks = []
    current_chunk = []
    current_pages = set()

    # Process decisions page by page, sorted by page number (None last)
    sorted_pages = sorted(by_page.keys(), key=lambda p: (p is None, p or 0))

    for page in sorted_pages:
        page_decisions = by_page[page]

        # If adding all decisions from this page would exceed limit, flush first
        if current_chunk and len(current_chunk) + len(page_decisions) > max_chunk_size:
            chunks.append(current_chunk)
            current_chunk = []
            current_pages = set()

        # Add decisions from this page (may need to split if too many on one page)
        for d in page_decisions:
            current_chunk.append(d)
            if page is not None:
                current_pages.add(page)

            if len(current_chunk) >= max_chunk_size:
                chunks.append(current_chunk)
                current_chunk = []
                current_pages = set()

    # Don't forget the last chunk
    if current_chunk:
        chunks.append(current_chunk)

    return chunks


# -------------------------
# VALIDATION RULES (lightweight)
# -------------------------
def rule_validation(final_obj):
    issues = []
    inc = (((final_obj.get("record") or {}).get("sheets")) or {}).get("included_articles") or {}
    if not isinstance(inc, dict):
        return issues

    # MRONJ development must be Yes/No/null (template validation).
    mdev = inc.get("mronj_development")
    if mdev not in (None, "Yes", "No"):
        issues.append({"severity": "WARN", "code": "MRONJ_DEV_UNEXPECTED", "message": "mronj_development should be Yes/No/blank to match template.", "path": "/record/sheets/included_articles/mronj_development"})

    # route_not_reported conflicts.
    route_nr = inc.get("route_not_reported") == 1
    route_any = any(inc.get(k) == 1 for k in ("route_iv", "route_oral", "route_im", "route_subcutaneous", "route_both"))
    if route_nr and route_any:
        issues.append({"severity": "WARN", "code": "ROUTE_NR_CONFLICT", "message": "route_not_reported is set but other route flags are also set.", "path": "/record/sheets/included_articles"})

    return issues


# -------------------------
# VERIFIER (reviews only non-null decisions)
# -------------------------
def gemini_verify_chunk(gclient, view_text, driver_json, decisions_to_review):
    user_text = (
        "PAPER_TEXT (VIEW):\n"
        + view_text
        + "\n\nDRIVER_JSON (context):\n"
        + json.dumps(driver_json, ensure_ascii=True)
        + "\n\nDECISIONS_TO_REVIEW:\n"
        + json.dumps(decisions_to_review, ensure_ascii=True)
    )
    return gemini_json_verifier(gclient, VERIFIER_SYSTEM, user_text, VERIFIER_SCHEMA)

def openai_verify_chunk(oai_client, view_text, driver_json, decisions_to_review):
    user_text = (
        "PAPER_TEXT (VIEW):\n"
        + view_text
        + "\n\nDRIVER_JSON (context):\n"
        + json.dumps(driver_json, ensure_ascii=True)
        + "\n\nDECISIONS_TO_REVIEW:\n"
        + json.dumps(decisions_to_review, ensure_ascii=True)
    )
    return openai_json_verifier(oai_client, VERIFIER_SYSTEM, user_text, VERIFIER_SCHEMA, "mronj_verifier_v2")


# -------------------------
# WORD REPORT
# -------------------------
def write_review_docx(final_obj, docx_path, append=True):
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    if append and os.path.exists(docx_path):
        doc = Document(docx_path)
        doc.add_page_break()
    else:
        doc = Document()
        doc.add_heading("MRONJ Prevention Extraction - Review Log", level=0)

    paper_id = final_obj.get("paper_id") or {}
    pmid = paper_id.get("pmid")
    doi = paper_id.get("doi")
    title = paper_id.get("title")

    # Paper header section
    doc.add_heading(f"PMID: {pmid if pmid is not None else 'N/A'}", level=1)
    if title:
        p = doc.add_paragraph()
        p.add_run("Title: ").bold = True
        p.add_run(str(title))
    if doi:
        p = doc.add_paragraph()
        p.add_run("DOI: ").bold = True
        p.add_run(str(doi))

    p = doc.add_paragraph()
    p.add_run("Study Type: ").bold = True
    p.add_run(str(final_obj.get("study_type", "N/A")))

    needs = ((final_obj.get("validation") or {}).get("needs_human_review"))
    p = doc.add_paragraph()
    p.add_run("Needs Human Review: ").bold = True
    status_run = p.add_run("YES" if needs else "NO")
    if needs:
        status_run.font.color.rgb = RGBColor(192, 0, 0)  # Red for YES
        status_run.bold = True
    else:
        status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for NO

    # Verifier decisions table
    doc.add_heading("Extraction Decisions", level=2)

    decisions = (final_obj.get("verification") or {}).get("critical_decisions") or []
    if not decisions:
        doc.add_paragraph("No decisions recorded.")
    else:
        # Create table with 5 columns: Field, Value, Status, Explanation, Evidence
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row styling
        hdr = t.rows[0].cells
        headers = ["Field", "Value", "Status", "Explanation", "Evidence"]
        for i, (cell, header_text) in enumerate(zip(hdr, headers)):
            cell.text = header_text
            # Bold header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
            # Light blue background
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E1F2"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for cd in decisions:
            row = t.add_row().cells
            path = str(cd.get("path", ""))
            # Extract just the field name from path like /record/sheets/included_articles/n_pts
            field_name = path.split("/")[-1] if "/" in path else path
            display_name = COLUMN_DISPLAY_NAMES.get(field_name, field_name.replace("_", " ").title())

            row[0].text = display_name
            row[1].text = str(cd.get("final_value", "")) if cd.get("final_value") is not None else ""
            row[2].text = str(cd.get("status", ""))
            row[3].text = str(cd.get("explanation", ""))
            row[4].text = str(cd.get("evidence", ""))

            # Color status cell
            status = cd.get("status", "")
            if status == "AGREE":
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="C6EFCE"/>')  # Light green
            elif status == "DISAGREE":
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFC7CE"/>')  # Light red
            else:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFEB9C"/>')  # Light yellow
            row[2]._tc.get_or_add_tcPr().append(shading)

            # Set font size for all cells
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Set column widths (approximate)
        for row in t.rows:
            row.cells[0].width = Inches(1.5)   # Field
            row.cells[1].width = Inches(1.2)   # Value
            row.cells[2].width = Inches(0.8)   # Status
            row.cells[3].width = Inches(2.5)   # Explanation
            row.cells[4].width = Inches(2.5)   # Evidence

    # Validation issues section
    doc.add_heading("Validation Issues", level=2)
    issues = (final_obj.get("validation") or {}).get("issues") or []
    if not issues:
        doc.add_paragraph("None.")
    else:
        for it in issues:
            severity = it.get("severity", "INFO")
            p = doc.add_paragraph()
            severity_run = p.add_run(f"[{severity}] ")
            severity_run.bold = True
            if severity == "CRITICAL":
                severity_run.font.color.rgb = RGBColor(192, 0, 0)
            elif severity == "WARN":
                severity_run.font.color.rgb = RGBColor(192, 128, 0)
            p.add_run(f"{it.get('code', '')}: {it.get('message', '')}")
            if it.get("path"):
                path_run = p.add_run(f" (path: {it.get('path')})")
                path_run.font.size = Pt(8)
                path_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(docx_path)


# -------------------------
# FINAL OBJECT BUILD
# -------------------------
def compile_critical_decision_report(verifier_passes, decisions_non_null, final_driver):
    issues = []
    latest_review_by_path = {}
    for p in verifier_passes or []:
        for dr in (p.get("decision_reviews") or []):
            path = dr.get("path")
            if path:
                latest_review_by_path[path] = dr

    critical_report = []
    for d in decisions_non_null:
        path = d.get("path")
        if not path:
            continue
        review = latest_review_by_path.get(path)
        if review is None:
            critical_report.append({
                "path": path,
                "final_value": json_pointer_get(final_driver, path),
                "status": "MISSING",
                "explanation": "Missing verifier review for decision.",
                "evidence": "",
            })
            issues.append({
                "severity": "CRITICAL",
                "code": "MISSING_VERIFIER_REVIEW",
                "message": f"Decision not reviewed by verifier: {path}",
                "path": path,
            })
            continue

        status = review.get("status", "UNSURE")
        final_val = review.get("proposed_value", review.get("driver_value"))

        critical_report.append({
            "path": path,
            "final_value": final_val,
            "status": status,
            "explanation": review.get("explanation", ""),
            "evidence": review.get("evidence", ""),
        })

        if status in ("DISAGREE", "UNSURE"):
            issues.append({
                "severity": "CRITICAL",
                "code": f"VERIFIER_{status}",
                "message": f"Verifier status {status} for decision: {path}",
                "path": path,
            })

    return critical_report, issues

def build_final_object(working_obj, verifier_passes, decisions_non_null, verifier_model):
    merged = copy.deepcopy(working_obj)
    # Apply all suggested patches.
    for p in verifier_passes or []:
        patch = p.get("suggested_patch")
        if isinstance(patch, dict) and patch:
            merged = deep_merge(merged, patch)

    compute_scores_inplace(merged)

    critical_report, issues = compile_critical_decision_report(verifier_passes, decisions_non_null, merged)
    issues.extend(rule_validation(merged))

    final_obj = {
        "version": "2.0",
        "paper_id": merged.get("paper_id") or {"pmid": None, "doi": None, "title": None},
        "study_type": merged.get("study_type", "unclear"),
        "record": merged.get("record") or {"sheets": {}},
        "verification": {
            "verifier_model": verifier_model,
            "passes": verifier_passes,
            "critical_decisions": critical_report,
        },
        "validation": {
            "needs_human_review": any(i.get("severity") == "CRITICAL" for i in issues),
            "issues": issues,
        },
    }
    return final_obj


# -------------------------
# TASK RUNNER (single paper)
# -------------------------
def _progress(progress_fn, message):
    if progress_fn:
        ts = datetime.now(UTC).strftime("%Y-%m-%d %H:%M:%S")
        progress_fn(f"[{ts} UTC] {message}")

def _init_working_object():
    return {
        "paper_id": {"pmid": None, "doi": None, "title": None},
        "study_type": "unclear",
        "record": {
            "sheets": {
                "included_articles": {},
                "level_of_evidence": {},
                "rct_appraisal": None,
                "cohort_appraisal": None,
                "case_series_appraisal": None,
                "case_control_appraisal": None,
                "systematic_appraisal": None,
            }
        },
    }

def _apply_patch(working, patch):
    if not isinstance(patch, dict) or not patch:
        return working
    return deep_merge(working, patch)

def _prune_redundant_patch_fields(patch: Optional[dict]) -> Optional[dict]:
    if not isinstance(patch, dict) or not patch:
        return patch
    cleaned = copy.deepcopy(patch)
    record = cleaned.get("record")
    if isinstance(record, dict):
        sheets = record.get("sheets")
        if sheets is None or not isinstance(sheets, dict):
            record.pop("sheets", None)
        if not record:
            cleaned.pop("record", None)
    paper_id = cleaned.get("paper_id")
    if isinstance(paper_id, dict) and all(paper_id.get(k) in (None, "") for k in ("pmid", "doi", "title")):
        cleaned.pop("paper_id", None)
    if cleaned.get("study_type") in (None, ""):
        cleaned.pop("study_type", None)
    return cleaned

def _collect_decisions(all_task_results):
    return dedupe_decisions(all_task_results)


def _pubmed_esearch(term: str, api_key: Optional[str], email: Optional[str], timeout: int) -> list[str]:
    if not term:
        return []
    params = {
        "db": "pubmed",
        "retmode": "json",
        "term": term,
    }
    if api_key:
        params["api_key"] = api_key
    if email:
        params["email"] = email
    url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?" + urllib.parse.urlencode(params)
    with urllib.request.urlopen(url, timeout=timeout) as response:
        payload = json.loads(response.read().decode("utf-8"))
    return (payload.get("esearchresult") or {}).get("idlist") or []


def lookup_pmid_via_pubmed(title: Optional[str], doi: Optional[str]) -> Optional[str]:
    if not ENABLE_PUBMED_LOOKUP:
        return None
    if doi:
        ids = _pubmed_esearch(f"{doi}[DOI]", PUBMED_API_KEY, PUBMED_EMAIL, PUBMED_LOOKUP_TIMEOUT)
        if ids:
            return ids[0]
    if title:
        ids = _pubmed_esearch(f"{title}[Title]", PUBMED_API_KEY, PUBMED_EMAIL, PUBMED_LOOKUP_TIMEOUT)
        if ids:
            return ids[0]
    return None

def run_pipeline_for_pdf(
    pdf_path,
    oai_client,
    gclient,
    template_xlsx,
    out_xlsx,
    out_docx,
    progress_fn=print,
    use_gemini_driver=False,
    use_openai_verifier=False,
    clear_existing_data=False,
):
    _progress(progress_fn, f"Starting PDF: {pdf_path}")
    full_pages = extract_pdf_pages(pdf_path)
    paper_id_hint = extract_paper_id_from_pages(full_pages)
    if paper_id_hint.get("doi") or paper_id_hint.get("title"):
        _progress(
            progress_fn,
            "PDF metadata hint: "
            + f"doi={paper_id_hint.get('doi') or 'N/A'} "
            + f"title={paper_id_hint.get('title') or 'N/A'}",
        )

    verifier_fn = openai_verify_chunk if use_openai_verifier else gemini_verify_chunk
    verifier_model = OPENAI_VERIFIER_MODEL if use_openai_verifier else GEMINI_VERIFIER_MODEL

    def run_driver(task_name, schema, user_prompt, schema_name):
        """Run extraction for a task."""
        if use_gemini_driver:
            return gemini_json(gclient, TASK_SYSTEM, user_prompt, schema)
        else:
            return openai_json(oai_client, TASK_SYSTEM, user_prompt, schema, schema_name)

    def verify_decisions(task_result, working_snapshot):
        """Verify decisions from a single task result immediately."""
        decisions = decisions_only_non_null(task_result.get("decisions") or [])
        if not decisions:
            return []

        verifier_passes = []
        chunks = chunk_list(decisions, VERIFIER_CHUNK_SIZE)
        working_json = sanitize_for_model_input(working_snapshot)

        for ch in chunks:
            page_numbers = sorted({d.get("page") for d in ch if d.get("page") is not None})
            verifier_pages = _select_pages(full_pages, page_numbers)
            verifier_view = build_verifier_view(verifier_pages, ch)
            vpass = verifier_fn(
                oai_client if use_openai_verifier else gclient,
                verifier_view,
                working_json,
                ch,
            )
            verifier_passes.append(vpass)
        return verifier_passes

    def run_task_with_verify(task_num, task_name, view_keywords, allowed_keys, schema_name,
                             fields_text, sheet_key="included_articles", allowed_level_keys=None):
        """Run a single task: extract + verify immediately. Returns (task_result, verifier_passes, patch)."""
        _progress(progress_fn, f"Task {task_num}: {task_name} starting...")

        view = make_task_view(full_pages, keywords=view_keywords)
        if allowed_level_keys:
            schema = build_task_schema(task_name=task_name, allowed_sheet_key=None,
                                       allowed_included_keys=allowed_keys, allowed_level_keys=allowed_level_keys)
        else:
            schema = build_task_schema(task_name=task_name, allowed_sheet_key=sheet_key,
                                       allowed_included_keys=allowed_keys)

        user_prompt = _task_user(task_name, fields_text, view, context_json=None)
        task_result = run_driver(task_name, schema, user_prompt, schema_name)

        _progress(progress_fn, f"Task {task_num}: {task_name} extracted, verifying...")

        # Build a minimal working object for verification context
        working_snapshot = _init_working_object()
        pruned_patch = _prune_redundant_patch_fields(task_result.get("patch"))
        working_snapshot = _apply_patch(working_snapshot, pruned_patch)

        verifier_passes = verify_decisions(task_result, working_snapshot)
        _progress(progress_fn, f"Task {task_num}: {task_name} done.")

        task_result["patch"] = pruned_patch
        return (task_result, verifier_passes, pruned_patch)

    # ---- Define all tasks ----
    task1_config = {
        "task_num": "1/5", "task_name": "meta_design",
        "view_keywords": ["pmid", "doi", "random", "cohort", "case", "systematic review", "methods", "abstract", "level of evidence", "grade", "recommendation", "oxford", "sign", "grade"],
        "allowed_keys": ["author", "year", "study_design"],
        "allowed_level_keys": ["level_of_evidence", "grade_of_recommendation"],
        "schema_name": "mronj_task_meta_design",
        "fields_text": (
            "EXTRACT these fields:\n"
            "- paper_id: doi, title (from paper header/abstract); leave pmid null (resolved via PubMed)\n"
            "- study_type: MUST be one of: " + "|".join(STUDY_TYPE_ENUM) + "\n"
            "- included_articles.author: first author surname (e.g. 'Smith')\n"
            "- included_articles.year: publication year as integer\n"
            "- included_articles.study_design: brief description (e.g. 'Retrospective cohort')\n"
            "- level_of_evidence.level_of_evidence: e.g. '1a', '2b', 'III' - only if explicitly stated\n"
            "- level_of_evidence.grade_of_recommendation: e.g. 'A', 'B', 'C' - only if explicitly stated"
        ),
    }

    task2_config = {
        "task_num": "2/5", "task_name": "population",
        "view_keywords": ["participants", "patients", "sample", "n=", "mean age", "male", "female", "table 1"],
        "allowed_keys": ["n_pts", "age_mean_years", "gender_male_n", "gender_female_n"],
        "schema_name": "mronj_task_population",
        "fields_text": (
            "EXTRACT these fields from included_articles sheet:\n"
            "- n_pts: total number of patients/participants as integer\n"
            "- age_mean_years: mean age in years as number (e.g. 65.4)\n"
            "- gender_male_n: number of male participants as integer\n"
            "- gender_female_n: number of female participants as integer\n"
            "NOTE: Leave paper_id and study_type as null (handled by another task)"
        ),
    }

    task3_config = {
        "task_num": "3/5", "task_name": "indication_drugs_route_site",
        "view_keywords": ["breast", "prostate", "myeloma", "osteoporosis", "zoled", "pamid", "alend", "rised", "iband", "etid", "clodron", "denos", "intraven", "oral", "subcut", "mandible", "maxilla"],
        "allowed_keys": [
            "site_maxilla","site_mandible","site_both",
            "primary_cause_breast_cancer","primary_cause_prostate_cancer","primary_cause_mm","primary_cause_osteoporosis","primary_cause_other",
            "ards_bisphosphonates_zoledronate","ards_bisphosphonates_pamidronate","ards_bisphosphonates_risedronate","ards_bisphosphonates_alendronate",
            "ards_bisphosphonates_ibandronate","ards_bisphosphonates_combination","ards_bisphosphonates_etidronate","ards_bisphosphonates_clodronate",
            "ards_bisphosphonates_unknown_other","ards_denosumab","ards_both",
            "route_iv","route_oral","route_im","route_subcutaneous","route_both","route_not_reported",
        ],
        "schema_name": "mronj_task_indication_drugs",
        "fields_text": (
            "EXTRACT these flags (1 if present, null if not mentioned) from included_articles sheet:\n"
            "SITE (where MRONJ occurred):\n"
            "- site_maxilla, site_mandible, site_both\n"
            "PRIMARY CAUSE/INDICATION:\n"
            "- primary_cause_breast_cancer, primary_cause_prostate_cancer, primary_cause_mm (multiple myeloma)\n"
            "- primary_cause_osteoporosis, primary_cause_other\n"
            "DRUGS - Bisphosphonates:\n"
            "- ards_bisphosphonates_zoledronate (Zometa/Reclast)\n"
            "- ards_bisphosphonates_pamidronate (Aredia)\n"
            "- ards_bisphosphonates_alendronate (Fosamax)\n"
            "- ards_bisphosphonates_risedronate (Actonel)\n"
            "- ards_bisphosphonates_ibandronate (Boniva)\n"
            "- ards_bisphosphonates_etidronate, ards_bisphosphonates_clodronate\n"
            "- ards_bisphosphonates_combination (multiple BPs), ards_bisphosphonates_unknown_other\n"
            "DRUGS - Other:\n"
            "- ards_denosumab (Prolia/Xgeva), ards_both (BP + denosumab)\n"
            "ROUTE of administration:\n"
            "- route_iv, route_oral, route_im, route_subcutaneous, route_both, route_not_reported\n"
            "NOTE: Leave paper_id and study_type as null (handled by another task)"
        ),
    }

    task4_config = {
        "task_num": "4/5", "task_name": "intervention_outcomes",
        "view_keywords": ["prevention", "dental", "extraction", "antibiotic", "photodynamic", "chlorhexidine", "follow-up", "months", "outcome", "mronj", "osteonecrosis"],
        "allowed_keys": [
            "mronj_stage_at_risk","mronj_stage_0",
            "prevention_technique","group_intervention","group_control",
            "follow_up_mean_months","follow_up_range","outcome_variable","mronj_development","mronj_development_details",
        ],
        "schema_name": "mronj_task_outcomes",
        "fields_text": (
            "EXTRACT these fields from included_articles sheet:\n"
            "STAGING:\n"
            "- mronj_stage_at_risk: number of patients at risk stage (integer)\n"
            "- mronj_stage_0: number of patients at stage 0 (integer)\n"
            "INTERVENTION:\n"
            "- prevention_technique: description of prevention method used\n"
            "- group_intervention: description of intervention group\n"
            "- group_control: description of control group\n"
            "FOLLOW-UP:\n"
            "- follow_up_mean_months: mean follow-up duration in months (number)\n"
            "- follow_up_range: follow-up range as string (e.g. '6-24 months')\n"
            "OUTCOMES:\n"
            "- outcome_variable: primary outcome measured\n"
            "- mronj_development: 'Yes' or 'No' - did MRONJ develop?\n"
            "- mronj_development_details: details about MRONJ cases if any\n"
            "NOTE: Leave paper_id and study_type as null (handled by another task)"
        ),
    }

    # ---- Run Tasks 1-4 in parallel (each with immediate verification) ----
    _progress(progress_fn, "Running tasks 1-4 in parallel...")

    all_task_results = []
    all_verifier_passes = []
    all_patches = []
    task1_result = None  # Need this for Task 5's study_type

    with ThreadPoolExecutor(max_workers=4) as executor:
        future_task1 = executor.submit(
            run_task_with_verify, task1_config["task_num"], task1_config["task_name"],
            task1_config["view_keywords"], task1_config["allowed_keys"], task1_config["schema_name"],
            task1_config["fields_text"], None, task1_config["allowed_level_keys"]
        )
        future_task2 = executor.submit(
            run_task_with_verify, task2_config["task_num"], task2_config["task_name"],
            task2_config["view_keywords"], task2_config["allowed_keys"], task2_config["schema_name"],
            task2_config["fields_text"]
        )
        future_task3 = executor.submit(
            run_task_with_verify, task3_config["task_num"], task3_config["task_name"],
            task3_config["view_keywords"], task3_config["allowed_keys"], task3_config["schema_name"],
            task3_config["fields_text"]
        )
        future_task4 = executor.submit(
            run_task_with_verify, task4_config["task_num"], task4_config["task_name"],
            task4_config["view_keywords"], task4_config["allowed_keys"], task4_config["schema_name"],
            task4_config["fields_text"]
        )

        # Collect results
        task1_result, task1_vpasses, task1_patch = future_task1.result()
        all_task_results.append(task1_result)
        all_verifier_passes.extend(task1_vpasses)
        all_patches.append(task1_patch)

        task2_result, task2_vpasses, task2_patch = future_task2.result()
        all_task_results.append(task2_result)
        all_verifier_passes.extend(task2_vpasses)
        all_patches.append(task2_patch)

        task3_result, task3_vpasses, task3_patch = future_task3.result()
        all_task_results.append(task3_result)
        all_verifier_passes.extend(task3_vpasses)
        all_patches.append(task3_patch)

        task4_result, task4_vpasses, task4_patch = future_task4.result()
        all_task_results.append(task4_result)
        all_verifier_passes.extend(task4_vpasses)
        all_patches.append(task4_patch)

    # ---- Task 5: Critical appraisal (needs study_type from Task 1) ----
    # Get study_type from task1 result
    task1_patch_obj = task1_patch or {}
    study_type = task1_patch_obj.get("study_type") or "unclear"

    if study_type in ("rct", "cohort", "case_series", "case_control", "systematic_review"):
        _progress(progress_fn, f"Task 5/5: critical appraisal ({study_type}) starting...")
        view5 = make_task_view(full_pages, keywords=["methods", "random", "blind", "withdraw", "confound", "follow up", "loss to follow up", "search strategy", "protocol", "meta-analysis", "risk of bias"])
        schema5 = build_appraisal_schema(study_type)
        fields_text = "- Fill only the appraisal sheet for study_type=" + study_type
        user5 = _task_user("critical_appraisal", fields_text, view5, context_json=None)

        task5_result = run_driver("critical_appraisal", schema5, user5, "mronj_task_appraisal")
        _progress(progress_fn, f"Task 5/5: critical appraisal extracted, verifying...")

        working5 = _init_working_object()
        working5 = _apply_patch(working5, task5_result.get("patch"))
        task5_vpasses = verify_decisions(task5_result, working5)

        all_task_results.append(task5_result)
        all_verifier_passes.extend(task5_vpasses)
        all_patches.append(task5_result.get("patch"))
        _progress(progress_fn, "Task 5/5: critical appraisal done.")
    else:
        _progress(progress_fn, "Task 5/5 skipped (study_type unclear/other).")

    # ---- Merge all results ----
    _progress(progress_fn, "Merging all task results...")
    working = _init_working_object()
    for patch in all_patches:
        if patch:
            working = _apply_patch(working, patch)

    # Apply verifier corrections
    for vpass in all_verifier_passes:
        if vpass:
            patch = vpass.get("suggested_patch")
            if isinstance(patch, dict) and patch:
                working = deep_merge(working, patch)

    if ENABLE_PUBMED_LOOKUP:
        working.setdefault("paper_id", {})["pmid"] = None

    working.setdefault("paper_id", {})
    if paper_id_hint.get("doi"):
        working["paper_id"]["doi"] = paper_id_hint.get("doi")
    if paper_id_hint.get("title"):
        working["paper_id"]["title"] = paper_id_hint.get("title")

    # Optional PMID lookup via PubMed if missing, using PDF-derived DOI/title.
    if ENABLE_PUBMED_LOOKUP and not (working.get("paper_id") or {}).get("pmid"):
        lookup_doi = paper_id_hint.get("doi")
        lookup_title = paper_id_hint.get("title")
        if lookup_doi or lookup_title:
            try:
                found_pmid = lookup_pmid_via_pubmed(lookup_title, lookup_doi)
            except Exception as exc:
                _progress(progress_fn, f"PubMed lookup failed: {exc}")
                found_pmid = None
            if found_pmid:
                working.setdefault("paper_id", {})["pmid"] = found_pmid
        else:
            _progress(progress_fn, "PubMed lookup skipped (no PDF DOI/title found).")

    # Ensure pmid is copied to included_articles + level_of_evidence for convenience.
    pmid = (working.get("paper_id") or {}).get("pmid")

    # Ensure nested structures exist
    if "record" not in working:
        working["record"] = {}
    if not isinstance(working["record"].get("sheets"), dict):
        working["record"]["sheets"] = {}
    if working["record"]["sheets"].get("included_articles") is None:
        working["record"]["sheets"]["included_articles"] = {}
    if working["record"]["sheets"].get("level_of_evidence") is None:
        working["record"]["sheets"]["level_of_evidence"] = {}

    if pmid is not None:
        working["record"]["sheets"]["included_articles"]["pmid"] = pmid
        working["record"]["sheets"]["level_of_evidence"]["pmid"] = pmid

    # Copy author/year/design if we have them.
    inc = working["record"]["sheets"].get("included_articles") or {}
    for k in ("author", "year", "study_design"):
        if inc.get(k) not in (None, ""):
            working["record"]["sheets"]["level_of_evidence"][k] = inc.get(k)

    # Collect all decisions for final object
    all_decisions = _collect_decisions(all_task_results)
    if ENABLE_PUBMED_LOOKUP:
        all_decisions = [
            d for d in all_decisions
            if d.get("path") not in (
                "/paper_id/pmid",
                "/record/sheets/included_articles/pmid",
                "/record/sheets/level_of_evidence/pmid",
            )
        ]
    decisions_non_null = decisions_only_non_null(all_decisions)

    final_obj = build_final_object(working, all_verifier_passes, decisions_non_null, verifier_model=verifier_model)

    _progress(progress_fn, "Writing Excel + Word outputs...")
    apply_to_workbook(final_obj, template_xlsx, out_xlsx, EXCEL_MAP, clear_existing_data=clear_existing_data)
    write_review_docx(final_obj, out_docx, append=True)

    audit_path = out_xlsx.replace(".xlsx", f".audit_{pmid}.json")
    with open(audit_path, "w", encoding="utf-8") as f:
        json.dump(final_obj, f, ensure_ascii=False, indent=2)

    _progress(progress_fn, f"Completed PDF: {pdf_path}")
    return final_obj


def run_pipeline(
    pdf_paths=None,
    template_xlsx=None,
    out_xlsx=OUT_XLSX,
    out_docx=OUT_DOCX,
    openai_api_key=None,
    google_api_key=None,
    progress_fn=print,
    use_gemini_driver=False,
    use_openai_verifier=False,
    skip_existing_evals=True,
    processed_state_path=None,
):
    if not pdf_paths:
        raise RuntimeError("pdf_paths is empty. Provide at least one PDF path.")

    # Use provided template or generate fresh one
    actual_template = get_or_create_template(template_xlsx, EXCEL_MAP)
    _progress(progress_fn, f"Using template: {actual_template}")

    openai_key = openai_api_key or os.getenv("OPENAI_API_KEY")
    google_key = google_api_key or os.getenv("GOOGLE_API_KEY")
    if not openai_key:
        raise RuntimeError("Missing OPENAI_API_KEY.")
    if not google_key:
        raise RuntimeError("Missing GOOGLE_API_KEY.")

    oai_client = OpenAI(api_key=openai_key)
    gclient = genai.Client(api_key=google_key)

    def load_processed_state(state_path):
        if not state_path or not os.path.exists(state_path):
            return []
        try:
            with open(state_path, "r", encoding="utf-8") as f:
                payload = json.load(f)
            if isinstance(payload, list):
                return payload
        except Exception as exc:
            _progress(progress_fn, f"Warning: failed to read processed state ({state_path}): {exc}")
        return []

    def write_processed_state(state_path, entries):
        if not state_path:
            return
        os.makedirs(os.path.dirname(state_path) or ".", exist_ok=True)
        with open(state_path, "w", encoding="utf-8") as f:
            json.dump(entries, f, ensure_ascii=False, indent=2)

    state_path = processed_state_path or f"{out_xlsx}.processed.json"
    processed_entries = load_processed_state(state_path)
    processed_set = {os.path.abspath(e.get("pdf_path")) for e in processed_entries if isinstance(e, dict) and e.get("pdf_path")}

    current_template = out_xlsx if os.path.exists(out_xlsx) else actual_template
    finals = []
    processed_this_run = 0
    for idx, pdf in enumerate(pdf_paths):
        if not os.path.exists(pdf):
            raise FileNotFoundError(pdf)

        abs_pdf = os.path.abspath(pdf)
        if skip_existing_evals and abs_pdf in processed_set:
            _progress(progress_fn, f"Skipping already processed PDF: {pdf}")
            continue

        # Clear old/demo data only when creating a new output workbook
        is_first_pdf = (processed_this_run == 0 and not os.path.exists(out_xlsx))

        final_obj = run_pipeline_for_pdf(
            pdf_path=pdf,
            oai_client=oai_client,
            gclient=gclient,
            template_xlsx=current_template,
            out_xlsx=out_xlsx,
            out_docx=out_docx,
            progress_fn=progress_fn,
            use_gemini_driver=use_gemini_driver,
            use_openai_verifier=use_openai_verifier,
            clear_existing_data=is_first_pdf,
        )
        current_template = out_xlsx
        finals.append(final_obj)
        processed_this_run += 1

        pid = final_obj.get("paper_id") or {}
        _progress(progress_fn, "DONE pdf=" + str(pdf) + " pmid=" + str(pid.get("pmid")) + " study_type=" + str(final_obj.get("study_type")) + " needs_human_review=" + str((final_obj.get("validation") or {}).get("needs_human_review")))

        processed_entry = {
            "pdf_path": abs_pdf,
            "pmid": pid.get("pmid"),
            "processed_at": datetime.now(UTC).isoformat(),
        }
        processed_entries.append(processed_entry)
        processed_set.add(abs_pdf)
        write_processed_state(state_path, processed_entries)

    _progress(progress_fn, f"WROTE_XLSX: {out_xlsx}")
    _progress(progress_fn, f"WROTE_DOCX: {out_docx}")
    return finals
